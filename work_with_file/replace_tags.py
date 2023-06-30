from docx import Document
from docx.shared import Pt
def replace_tags(Report_name, tags):  # на вход принемается номе отчёты и список пар {тэг : текст для замены}
    doc = Document(Report_name)
    # идём по парвграфам и заменяем тэги на соответствующие им значения
    for p in doc.paragraphs:
        # идём по списку тэгов и зименяем их
        for elem in tags:
            # elem[0] - тэг  elem[1] - текст соответствующи тэгу
            if elem[0] in p.text:
                style = p.style
                text = p.text.replace(elem[0], elem[1])
                p.text = text
                p.style = style
                # замена стиля заголовка
                if elem[0] == "${Number}":
                    # запоминаем текст в параграфе
                    text = p.text
                    # очищаем параграф
                    p.text = ""
                    # вставляем тест с нужным стилем
                    runner = p.add_run(text)
                    runner.bold = True
                    font = runner.font
                    font.size = Pt(16)

                # замена стиля списка материалов
                if elem[0] == "${Video}" or elem[0] == "${Implant}":
                    # запоминаем текст в параграфе
                    text = p.text
                    # очищаем параграф
                    p.text = ""
                    # вставляем тест с нужным стилем
                    runner = p.add_run(text)
                    runner.italic = True



    doc.save(Report_name)


'''                    runner = p.add_run(elem[1])
                    runner.bold = True
                    font = runner.font
                    font.size = Pt(16)'''